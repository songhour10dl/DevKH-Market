
import csv
import os
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Inches
from data_models import JobListing

class ExportManager:
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)
    
    def export_to_csv(self, jobs: List[JobListing], skill_stats: Dict = None) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"job_analysis_{timestamp}.csv"
        filepath = os.path.join(self.export_dir, filename)
        
        try:
            with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
                fieldnames = [
                    'Title', 'Company', 'Location', 'Source Site', 
                    'URL', 'Identified Skills', 'Skills Count', 
                    'Description Preview', 'Scraped At'
                ]
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                
                for job in jobs:
                    writer.writerow({
                        'Title': job.title,
                        'Company': job.company,
                        'Location': job.location,
                        'Source Site': job.source_site,
                        'URL': job.url,
                        'Identified Skills': ', '.join(job.identified_skills),
                        'Skills Count': len(job.identified_skills),
                        'Description Preview': job.description[:200] + '...' if len(job.description) > 200 else job.description,
                        'Scraped At': job.scraped_at.strftime("%Y-%m-%d %H:%M:%S")
                    })
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Error exporting to CSV: {str(e)}")
    
    def export_to_word(self, jobs: List[JobListing], skill_stats: Dict = None) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"job_analysis_report_{timestamp}.docx"
        filepath = os.path.join(self.export_dir, filename)
        
        try:
            doc = Document()
            
            title = doc.add_heading('Cambodian Software Engineering Job Market Analysis', 0)
            
            doc.add_heading('Executive Summary', level=1)
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Analysis Date: ").bold = True
            summary_para.add_run(datetime.now().strftime("%B %d, %Y"))
            summary_para.add_run(f"\nTotal Jobs Analyzed: ").bold = True
            summary_para.add_run(str(len(jobs)))
            
            if skill_stats:
                summary_para.add_run(f"\nUnique Skills Identified: ").bold = True
                summary_para.add_run(str(skill_stats.get('unique_skills_found', 0)))
            
            if skill_stats and 'most_demanded_skills' in skill_stats:
                doc.add_heading('Most In-Demand Skills', level=1)
                skills_table = doc.add_table(rows=1, cols=2)
                skills_table.style = 'Table Grid'
                hdr_cells = skills_table.rows[0].cells
                hdr_cells[0].text = 'Skill'
                hdr_cells[1].text = 'Frequency'
                
                for skill, count in skill_stats['most_demanded_skills']:
                    row_cells = skills_table.add_row().cells
                    row_cells[0].text = skill
                    row_cells[1].text = str(count)
            
            if skill_stats and 'category_breakdown' in skill_stats:
                doc.add_heading('Skills by Category', level=1)
                for category, data in skill_stats['category_breakdown'].items():
                    doc.add_heading(category, level=2)
                    cat_para = doc.add_paragraph()
                    cat_para.add_run(f"Total Mentions: ").bold = True
                    cat_para.add_run(str(data['total_mentions']))
                    cat_para.add_run(f"\nUnique Skills: ").bold = True
                    cat_para.add_run(str(data['unique_skills']))
                    
                    if data['top_skills']:
                        cat_para.add_run(f"\nTop Skills: ").bold = True
                        top_skills_text = ', '.join([f"{skill} ({count})" for skill, count in data['top_skills']])
                        cat_para.add_run(top_skills_text)
            
            doc.add_heading('Job Listings Details', level=1)
            
            for i, job in enumerate(jobs[:50], 1):
                doc.add_heading(f"{i}. {job.title}", level=2)
                
                job_para = doc.add_paragraph()
                job_para.add_run("Company: ").bold = True
                job_para.add_run(job.company)
                job_para.add_run("\nLocation: ").bold = True
                job_para.add_run(job.location)
                job_para.add_run("\nSource: ").bold = True
                job_para.add_run(job.source_site)
                
                if job.identified_skills:
                    job_para.add_run("\nIdentified Skills: ").bold = True
                    job_para.add_run(', '.join(job.identified_skills))
                
                if job.url:
                    job_para.add_run("\nURL: ").bold = True
                    job_para.add_run(job.url)
                
                if job.description:
                    desc_preview = job.description[:300] + '...' if len(job.description) > 300 else job.description
                    job_para.add_run("\nDescription: ").bold = True
                    job_para.add_run(desc_preview)
                
                doc.add_paragraph()
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Error exporting to Word: {str(e)}")
    
    def get_export_directory(self) -> str:
        return os.path.abspath(self.export_dir)
